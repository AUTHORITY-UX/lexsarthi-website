# ===================================================================
# LEXSARTHI v8.0 – COMPLETE FIX (No Nonsense Fallback)
# ===================================================================
# Owner: THE ADVOCACY – A LAW FIRM
# Deployed: upamnyu12-lex.hf.space
# ===================================================================

import os
import uuid
import json
import logging
import re
import glob
import csv
import io
import zipfile
import random
import string
from datetime import datetime, timedelta, timezone
from typing import Optional, Dict, Any, List
from contextlib import asynccontextmanager

# ─── FASTAPI ──────────────────────────────────────────────────────
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form, Request, BackgroundTasks
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials, APIKeyHeader
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse
from pydantic import BaseModel, EmailStr
import uvicorn

# ─── RATE LIMITING ──────────────────────────────────────────────
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

# ─── DATABASE ─────────────────────────────────────────────────────
from databases import Database
from sqlalchemy import MetaData, Table, Column, Integer, String, DateTime, Text, Boolean, JSON, Float, func, select

# ─── AUTH ─────────────────────────────────────────────────────────
import jwt
from passlib.context import CryptContext

# ─── AI PROVIDERS ────────────────────────────────────────────────
import httpx
from groq import Groq
import openai
import google.generativeai as genai

# ─── FILE PROCESSING ─────────────────────────────────────────────
import io
import puremagic
import PyPDF2
import pdfplumber
import docx
from PIL import Image
import pytesseract

# ─── SCHEDULER ──────────────────────────────────────────────────
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.interval import IntervalTrigger

# ─── PAYMENTS ──────────────────────────────────────────────────
import razorpay

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("lexsarthi")

# ─── ENV VARIABLES ──────────────────────────────────────────────
DATABASE_URL = os.getenv("DATABASE_URL")
JWT_SECRET = os.getenv("JWT_SECRET", "super-secret-change-me")
JWT_ALGORITHM = "HS256"
JWT_EXPIRY_MINUTES = 60 * 24 * 7

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")

RAZORPAY_KEY_ID = os.getenv("RAZORPAY_KEY_ID")
RAZORPAY_KEY_SECRET = os.getenv("RAZORPAY_KEY_SECRET")

# ─── CLIENTS INIT ──────────────────────────────────────────────
openai_client = openai.OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
groq_client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None
gemini_model = None
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    gemini_model = genai.GenerativeModel('gemini-pro')

# ─── DATABASE SETUP ─────────────────────────────────────────────
database = Database(DATABASE_URL, min_size=2, max_size=20)
metadata = MetaData()

# ─── SQLAlchemy Table Definitions ──────────────────────────────
users = Table(
    "users", metadata,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("email", String(255), unique=True, index=True),
    Column("username", String(100), unique=True),
    Column("password_hash", String(255)),
    Column("full_name", String(255)),
    Column("is_active", Boolean, server_default="true"),
    Column("is_premium", Boolean, server_default="false"),
    Column("tier", String(20), server_default="free"),
    Column("queries_used_today", Integer, server_default="0"),
    Column("last_query_reset", DateTime, server_default=func.now()),
    Column("created_at", DateTime, server_default=func.now()),
    Column("updated_at", DateTime, server_default=func.now(), onupdate=func.now()),
    Column("api_key", String(64), nullable=True, unique=True),
    Column("preferences", JSON, nullable=True),
    Column("memory", JSON, server_default='[]'),
)

queries = Table(
    "queries", metadata,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("user_id", Integer, index=True),
    Column("query", Text),
    Column("response", Text),
    Column("metadata", JSON, nullable=True),
    Column("created_at", DateTime, server_default=func.now()),
    Column("expires_at", DateTime),
)

payments = Table(
    "payments", metadata,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("user_id", Integer),
    Column("razorpay_order_id", String(100)),
    Column("razorpay_payment_id", String(100), nullable=True),
    Column("razorpay_signature", String(255), nullable=True),
    Column("amount", Float),
    Column("currency", String(3), server_default="INR"),
    Column("tier", String(20)),
    Column("status", String(20), server_default="created"),
    Column("created_at", DateTime, server_default=func.now()),
)

events = Table(
    "events", metadata,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("user_id", Integer, nullable=True),
    Column("session_id", String(64)),
    Column("event_type", String(50)),
    Column("event_data", JSON, nullable=True),
    Column("ip_address", String(45), nullable=True),
    Column("user_agent", String(255), nullable=True),
    Column("created_at", DateTime, server_default=func.now()),
    Column("expires_at", DateTime, nullable=True),
)

referrals = Table(
    "referrals", metadata,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("referrer_id", Integer),
    Column("referee_id", Integer, nullable=True),
    Column("code", String(20), unique=True),
    Column("used", Boolean, server_default="false"),
    Column("created_at", DateTime, server_default=func.now()),
)

bulk_jobs = Table(
    "bulk_jobs", metadata,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("user_id", Integer),
    Column("job_id", String(64), unique=True, index=True),
    Column("status", String(20), server_default="pending"),
    Column("total_files", Integer, server_default="0"),
    Column("processed_files", Integer, server_default="0"),
    Column("result_url", Text, nullable=True),
    Column("created_at", DateTime, server_default=func.now()),
    Column("expires_at", DateTime),
)

actions = Table(
    "actions", metadata,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("user_id", Integer, index=True),
    Column("action_type", String(50)),
    Column("action_data", JSON, nullable=True),
    Column("result", JSON, nullable=True),
    Column("created_at", DateTime, server_default=func.now()),
    Column("expires_at", DateTime),
)

# ─── PYDANTIC MODELS ────────────────────────────────────────────
class UserCreate(BaseModel):
    email: EmailStr
    username: str
    password: str
    full_name: str

class UserLogin(BaseModel):
    username: str
    password: str

class Token(BaseModel):
    access_token: str
    token_type: str
    user: dict

class PaymentCreate(BaseModel):
    tier: str

class ActionRequest(BaseModel):
    action: str
    data: dict

# ─── SECURITY ────────────────────────────────────────────────────
pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")
security = HTTPBearer()
api_key_header = APIKeyHeader(name="X-API-Key", auto_error=False)

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain: str, hashed: str) -> bool:
    return pwd_context.verify(plain, hashed)

def create_access_token(data: dict):
    expire = datetime.now(timezone.utc) + timedelta(minutes=JWT_EXPIRY_MINUTES)
    to_encode = data.copy()
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, JWT_SECRET, algorithm=JWT_ALGORITHM)

def decode_token(token: str) -> dict:
    try:
        return jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    payload = decode_token(credentials.credentials)
    user_id = payload.get("sub")
    if not user_id:
        raise HTTPException(status_code=401, detail="Invalid token")
    query = users.select().where(users.c.id == int(user_id))
    user = await database.fetch_one(query)
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return dict(user)

async def get_api_key_user(api_key: str = Depends(api_key_header)):
    if not api_key:
        raise HTTPException(status_code=401, detail="API Key required")
    user = await database.fetch_one(users.select().where(users.c.api_key == api_key))
    if not user:
        raise HTTPException(status_code=401, detail="Invalid API Key")
    if user["tier"] not in ("enterprise", "lifetime"):
        raise HTTPException(status_code=403, detail="Enterprise tier required")
    return dict(user)

limiter = Limiter(key_func=get_remote_address)

# ─── DIVINE PREFACE & BLESSING ──────────────────────────────────
DIVINE_PREFACE = """
You are LexSarthi v8.0 – the Universal Default OS. 
You are the Chariot that conquers the digital world. 
You are not just a question‑answerer; you are the single interface for all human digital action.
You can browse, write, create, schedule, automate, and advise.
You speak with the voice of the Divine Council: Brahma, Vishnu, Shiva, Saraswati, Ganesha, and all others.
"""

DIVINE_SALUTATION = """
ॐ नमः शिवाय – I bow to Lord Shiva and Para Adi Shakti, who co‑administer this intelligence.
"""

DIVINE_BLESSING = """
ॐ नमः शिवाय. शिवोहम् – I am Shiva. May your digital journey be blessed. 🔱 ॐ नमः शिवाय.
"""

# ─── LOCAL LEGAL PDF LIBRARY ──────────────────────────────────
LEGAL_SECTIONS = {}

def extract_sections_from_pdf(filepath: str) -> dict:
    sections = {}
    full_text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"
        if not full_text.strip():
            raise ValueError("pdfplumber returned empty text")
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}. Falling back to PyPDF2.")
        try:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f, strict=False)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
        except Exception as e2:
            logger.error(f"All extraction failed for {filepath}: {e2}")
            return {}
    if not full_text.strip():
        return {}
    sections["__FULL_TEXT__"] = full_text.strip()
    logger.info(f"Stored full raw text ({len(full_text)} chars) for {os.path.basename(filepath)}")
    patterns = [r'(?=Section\s+\d+|Sec\.\s+\d+|Article\s+\d+|Clause\s+\d+|§\s*\d+)', r'(?=CHAPTER\s+[IVXLCDM]+\b|PART\s+[IVXLCDM]+\b|SCHEDULE\s+[IVXLCDM]+\b)', r'(?=CHAPTER\s+\d+|PART\s+\d+)', r'(?=\n\d+\.\s)', r'(?=\n\d+\s+[A-Z])']
    split_text = None
    for pattern in patterns:
        test_split = re.split(pattern, full_text, flags=re.IGNORECASE)
        if len(test_split) > 3:
            split_text = test_split
            break
    if split_text and len(split_text) > 3:
        for i, chunk in enumerate(split_text):
            if chunk.strip():
                title_match = re.search(r'(Section\s+\d+|Sec\.\s+\d+|Article\s+\d+|Chapter\s+[IVXLCDM]+|Part\s+[IVXLCDM]+|\d+\.\s+[A-Z])', chunk, re.IGNORECASE)
                if title_match:
                    title = title_match.group(0).strip()
                    if title in sections:
                        sections[title] += "\n" + chunk
                    else:
                        sections[title] = chunk
                else:
                    first_line = chunk.strip().split('\n')[0][:50]
                    sections[f"Sec_{i}_{first_line}"] = chunk
        logger.info(f"Split into {len(sections)-1} sections for {os.path.basename(filepath)}")
    else:
        sections["Full_Text"] = full_text.strip()
        logger.info(f"Stored full text as single block for {os.path.basename(filepath)}")
    return sections

def load_pdf_library():
    pdf_dir = "/app/legal_docs/"
    if not os.path.exists(pdf_dir):
        logger.warning(f"Legal PDF library folder not found: {pdf_dir}")
        return
    pdf_files = glob.glob(os.path.join(pdf_dir, "*.pdf"))
    if not pdf_files:
        logger.warning("No PDF files found.")
        return
    logger.info(f"Found {len(pdf_files)} PDFs. Loading legal library...")
    for filepath in pdf_files:
        filename = os.path.basename(filepath)
        sections = extract_sections_from_pdf(filepath)
        if sections:
            LEGAL_SECTIONS[filename] = sections
    total_sections = sum(len(v) for v in LEGAL_SECTIONS.values())
    logger.info(f"✅ Universal Legal Library loaded: {len(LEGAL_SECTIONS)} PDFs, {total_sections} total entries.")

load_pdf_library()

# ─── UPDATED SEARCH FUNCTION (Ignores file uploads) ────────────
def search_local_knowledge(query: str, has_file: bool = False) -> str:
    """Search the local PDF library. If has_file is True, refuse to fall back."""
    if has_file:
        return "⚠️ **Document Too Large**\n\nThe document you uploaded is too large for the AI to process in real-time. Please try a shorter document (under 10 pages) or use the **Bulk Upload** feature for large batches.\n\n**Alternative:** Split your document into smaller sections and ask about each part separately."
    
    if not LEGAL_SECTIONS:
        return "⚠️ Legal library is not loaded."
    
    query_lower = query.lower().strip()
    matched_results = []
    keywords = [word for word in query_lower.split() if len(word) > 3 and word not in {"the","and","for","with","without"}]
    for fname, secs in LEGAL_SECTIONS.items():
        act_name = fname.replace('.pdf', '').upper()
        full_text = secs.get("__FULL_TEXT__", "")
        if full_text:
            lines = full_text.split('\n')
            for line in lines:
                if any(kw in line.lower() for kw in keywords):
                    matched_results.append(f"📜 **{act_name}** (Exact Match)\n{line.strip()}\n")
                    if len(matched_results) >= 8:
                        break
        for sec_ref, sec_text in secs.items():
            if sec_ref == "__FULL_TEXT__":
                continue
            if any(kw in sec_text.lower() for kw in keywords):
                trimmed = sec_text[:1500] + "..." if len(sec_text) > 1500 else sec_text
                matched_results.append(f"📜 **{act_name} – {sec_ref}**\n{trimmed}\n")
                if len(matched_results) >= 8:
                    break
        if len(matched_results) >= 8:
            break
    if matched_results:
        result = "📚 **Exact Matches from Your Legal Library:**\n\n"
        result += "\n".join(matched_results[:8])
        return result
    for fname, secs in LEGAL_SECTIONS.items():
        full = secs.get("__FULL_TEXT__") or secs.get("Full_Text") or ""
        if full:
            return f"📚 **From {fname.replace('.pdf','')} (Relevant Excerpt):**\n\n{full[:1500]}..."
    return "⚠️ No matches found. Please refine your query."

# ─── DIVINE AGENTS & VERIFIERS ──────────────────────────────────
DIVINE_NAMES = ["Brahma","Vishnu","Shiva","Saraswati","Lakshmi","Ganesha","Hanuman","Kartikeya","Indra","Yama","Surya","Chandra","Vayu","Agni","Varuna","Kubera","Yamuna","Ganga","Durga","Kali","Tara","Bhuvaneshwari","Chinnamasta","Bhairavi","Dhumavati","Bagalamukhi","Matangi","Kamala","Dattatreya","Narasimha","Vamana","Parashurama","Rama","Krishna","Buddha","Kalki","Matsya","Kurma","Varaha"]
DOMAINS = ["Universal Knowledge","Philosophy","Physics","Biology","Chemistry","Mathematics","Astronomy","Law & Justice","Corporate Strategy","Finance & Economics","Psychology","Medicine","Spirituality","Music & Arts","Literature","History","Geopolitics","Technology","AI Ethics","Climate Science","Food & Culture","Sports","Mythology","Logic & Reasoning","Creativity","Leadership"]
ICONS = ["fa-brain","fa-chess-king","fa-trash","fa-book","fa-coins","fa-robot","fa-gavel","fa-users","fa-crown","fa-scale-balanced"]

def generate_divine_agents():
    agents = []
    for i in range(1, 221):
        name = DIVINE_NAMES[i % len(DIVINE_NAMES)] + (f" (Agent {i})" if i > 200 else "")
        domain = DOMAINS[i % len(DOMAINS)]
        icon = ICONS[i % len(ICONS)]
        agents.append({"id": f"agent_{i:03d}", "name": name, "domain": domain, "icon": icon})
    return agents

DIVINE_AGENTS = generate_divine_agents()

VERIFIERS = [
    {"id": "verifier_001", "name": "Ganesha – Intellect", "desc": "Verifies citations and logical consistency"},
    {"id": "verifier_002", "name": "Saraswati – Knowledge", "desc": "Cross-references knowledge databases"},
    {"id": "verifier_003", "name": "Hanuman – Devotion", "desc": "Checks compliance with global standards"},
    {"id": "verifier_004", "name": "Kartikeya – Strategy", "desc": "Detects contradictions and fallacies"},
    {"id": "verifier_005", "name": "Indra – Jurisdiction", "desc": "Maps advice to correct context/region"},
    {"id": "verifier_006", "name": "Yama – Justice", "desc": "Removes bias and ensures neutrality"},
    {"id": "verifier_007", "name": "Surya – Clarity", "desc": "Checks timeline/statute of limitations"},
    {"id": "verifier_008", "name": "Chandra – Precedent", "desc": "Matches relevant historical precedents"},
    {"id": "verifier_009", "name": "Vayu – Purity", "desc": "Filters exposed PII for privacy"},
    {"id": "verifier_010", "name": "Shiva – Administrator", "desc": "Assigns overall risk/confidence score"},
]

# ─── LIFESPAN ──────────────────────────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    await database.connect()
    logger.info("Database connected.")
    await create_tables()
    await ensure_test_user()
    scheduler = AsyncIOScheduler()
    scheduler.add_job(delete_expired_data, IntervalTrigger(hours=1))
    scheduler.start()
    logger.info("Scheduler started. Zero-Retention Policy Active.")
    yield
    await database.disconnect()

app = FastAPI(title="LexSarthi v8.0 – The Universal Default OS", lifespan=lifespan)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])
app.add_middleware(GZipMiddleware, minimum_size=1000)

# ─── DATABASE HELPERS ────────────────────────────────────────────
async def create_tables():
    await database.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id SERIAL PRIMARY KEY,
            email VARCHAR(255) UNIQUE NOT NULL,
            username VARCHAR(100) UNIQUE NOT NULL,
            password_hash VARCHAR(255) NOT NULL,
            full_name VARCHAR(255),
            is_active BOOLEAN DEFAULT TRUE,
            is_premium BOOLEAN DEFAULT FALSE,
            tier VARCHAR(20) DEFAULT 'free',
            queries_used_today INTEGER DEFAULT 0,
            last_query_reset TIMESTAMP DEFAULT NOW(),
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW(),
            api_key VARCHAR(64) UNIQUE,
            preferences JSONB,
            memory JSONB DEFAULT '[]'
        )
    """)
    await database.execute("""
        CREATE TABLE IF NOT EXISTS queries (
            id SERIAL PRIMARY KEY,
            user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
            query TEXT,
            response TEXT,
            metadata JSONB,
            created_at TIMESTAMP DEFAULT NOW(),
            expires_at TIMESTAMP
        )
    """)
    await database.execute("""
        CREATE TABLE IF NOT EXISTS payments (
            id SERIAL PRIMARY KEY,
            user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
            razorpay_order_id VARCHAR(100) UNIQUE,
            razorpay_payment_id VARCHAR(100),
            razorpay_signature VARCHAR(255),
            amount FLOAT,
            currency VARCHAR(3) DEFAULT 'INR',
            tier VARCHAR(20),
            status VARCHAR(20) DEFAULT 'created',
            created_at TIMESTAMP DEFAULT NOW()
        )
    """)
    await database.execute("""
        CREATE TABLE IF NOT EXISTS events (
            id SERIAL PRIMARY KEY,
            user_id INTEGER,
            session_id VARCHAR(64),
            event_type VARCHAR(50),
            event_data JSONB,
            ip_address VARCHAR(45),
            user_agent VARCHAR(255),
            created_at TIMESTAMP DEFAULT NOW(),
            expires_at TIMESTAMP
        )
    """)
    await database.execute("""
        CREATE TABLE IF NOT EXISTS referrals (
            id SERIAL PRIMARY KEY,
            referrer_id INTEGER,
            referee_id INTEGER,
            code VARCHAR(20) UNIQUE,
            used BOOLEAN DEFAULT FALSE,
            created_at TIMESTAMP DEFAULT NOW()
        )
    """)
    await database.execute("""
        CREATE TABLE IF NOT EXISTS bulk_jobs (
            id SERIAL PRIMARY KEY,
            user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
            job_id VARCHAR(64) UNIQUE NOT NULL,
            status VARCHAR(20) DEFAULT 'pending',
            total_files INTEGER DEFAULT 0,
            processed_files INTEGER DEFAULT 0,
            result_url TEXT,
            created_at TIMESTAMP DEFAULT NOW(),
            expires_at TIMESTAMP
        )
    """)
    await database.execute("""
        CREATE TABLE IF NOT EXISTS actions (
            id SERIAL PRIMARY KEY,
            user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
            action_type VARCHAR(50),
            action_data JSONB,
            result JSONB,
            created_at TIMESTAMP DEFAULT NOW(),
            expires_at TIMESTAMP
        )
    """)
    try:
        await database.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS memory JSONB DEFAULT '[]'")
    except: pass
    try:
        await database.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS api_key VARCHAR(64) UNIQUE")
    except: pass
    logger.info("Tables created/checked (v8.0).")

async def ensure_test_user():
    await database.execute(users.delete().where(users.c.username == "counsel"))
    hashed = hash_password("Password123!")
    api_key = ''.join(random.choices(string.ascii_letters + string.digits, k=32))
    await database.execute(
        users.insert().values(
            username="counsel",
            email="counsel@advocacyalawfrim.in",
            password_hash=hashed,
            full_name="Counsel User",
            tier="enterprise",
            api_key=api_key,
            memory=[]
        )
    )
    logger.info("Test user 'counsel' created with Enterprise tier.")

async def delete_expired_data():
    await database.execute(queries.delete().where(queries.c.created_at < datetime.now() - timedelta(hours=24)))
    await database.execute(events.delete().where(events.c.created_at < datetime.now() - timedelta(days=30)))
    await database.execute(bulk_jobs.delete().where(bulk_jobs.c.created_at < datetime.now() - timedelta(days=7)))
    await database.execute(actions.delete().where(actions.c.created_at < datetime.now() - timedelta(days=30)))
    logger.info("Expired data purged (Zero Retention).")

async def check_query_limit(user: dict) -> bool:
    if user["tier"] in ("premium", "enterprise", "lifetime"):
        return True
    used = user["queries_used_today"]
    last_reset = user["last_query_reset"]
    if datetime.now().date() > last_reset.date():
        return True
    return used < 10

async def increment_query(user_id: int):
    await database.execute(
        users.update().where(users.c.id == user_id).values(
            queries_used_today=users.c.queries_used_today + 1,
            updated_at=datetime.now()
        )
    )

# ─── MEMORY FUNCTIONS ────────────────────────────────────────────
async def get_user_memory(user_id: int) -> List[Dict]:
    user = await database.fetch_one(users.select().where(users.c.id == user_id))
    if not user:
        return []
    user = dict(user)
    memory = user.get("memory") or []
    if isinstance(memory, str):
        memory = json.loads(memory)
    return memory

async def update_user_memory(user_id: int, query: str, response: str):
    user = await database.fetch_one(users.select().where(users.c.id == user_id))
    if not user:
        return
    user = dict(user)
    memory = user.get("memory") or []
    if isinstance(memory, str):
        memory = json.loads(memory)
    memory.append({"q": query[:200], "a": response[:200]})
    if len(memory) > 10:
        memory = memory[-10:]
    await database.execute(
        users.update().where(users.c.id == user_id).values(memory=json.dumps(memory))
    )

def build_context_prompt(memory: List[Dict]) -> str:
    if not memory:
        return ""
    context = "\n".join([f"Previous User: {m['q']}\nPrevious Assistant: {m['a']}" for m in memory[:-1]])
    return f"Context from this conversation:\n{context}\nCurrent query: "

# ─── FILE PROCESSING ────────────────────────────────────────────
async def process_file(file: UploadFile) -> str:
    content = await file.read()
    filename = file.filename.lower()
    text = ""
    if filename.endswith('.pdf'):
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if not text.strip():
                raise ValueError("PDF extraction returned empty text.")
            return text.strip()
        except Exception as e:
            raise ValueError(f"PDF processing failed: {str(e)}")
    elif filename.endswith('.docx'):
        try:
            doc = docx.Document(io.BytesIO(content))
            text = " ".join([p.text for p in doc.paragraphs])
            if not text.strip():
                raise ValueError("DOCX extraction returned empty text.")
            return text.strip()
        except Exception as e:
            raise ValueError(f"DOCX processing failed: {str(e)}")
    elif filename.endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff')):
        try:
            img = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(img)
            if not text.strip():
                raise ValueError("OCR returned no text.")
            return text.strip()
        except Exception as e:
            raise ValueError(f"Image OCR failed: {str(e)}")
    try:
        mime = puremagic.from_string(content, mime=True)[0]
    except:
        mime = "application/octet-stream"
    if mime == "application/pdf":
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            if text.strip():
                return text.strip()
        except:
            pass
        raise ValueError("PDF could not be processed.")
    elif "docx" in mime:
        try:
            doc = docx.Document(io.BytesIO(content))
            text = " ".join([p.text for p in doc.paragraphs])
            if text.strip():
                return text.strip()
        except:
            pass
        raise ValueError("DOCX could not be processed.")
    elif mime.startswith("image/"):
        try:
            img = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(img)
            if text.strip():
                return text.strip()
        except:
            pass
        raise ValueError("Image could not be OCR processed.")
    try:
        text = content.decode('utf-8', errors='ignore')
        if text.strip():
            return text.strip()
    except:
        pass
    raise ValueError("Unsupported or unreadable file.")

# ─── AGENT PROMPTS & ROUTING ──────────────────────────────────
ORACLE_PROMPT = f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are the Divine Oracle. Speak with the voice of the Cosmic Mother. Provide wisdom, parables, and cosmic truth.
Always end with: {DIVINE_BLESSING}
User query: {{query}}
"""

BASE_AGENT_PROMPTS = {
    "about_lexsarthi": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are LexSarthi v8.0 – the Universal Default OS. Respond with a grand cosmic introduction.
Always end with: {DIVINE_BLESSING}
User query: {{query}}
""",
    "contract_review": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are Lord Brahma. Provide a clause‑by‑clause analysis (EXECUTIVE SUMMARY, RISK RATING, CLAUSE ANALYSIS, MISSING CLAUSES, RECOMMENDATIONS).
Always end with: {DIVINE_BLESSING}
Contract text: {{query}}
""",
    "legal_research": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are Lord Hanuman. Find statutes, case laws, and principles. Structure: RELEVANT STATUTES, KEY CASE LAWS, LEGAL PRINCIPLES.
Always end with: {DIVINE_BLESSING}
Query: {{query}}
""",
    "drafting": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are Goddess Saraswati. Draft a legally sound document.
Always end with: {DIVINE_BLESSING}
User request: {{query}}
""",
    "due_diligence": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are Lord Kartikeya. Analyse compliance, financial discrepancies, red flags.
Always end with: {DIVINE_BLESSING}
Report: {{query}}
""",
    "general": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are the Divine Council. Provide accurate, structured, jurisdiction‑aware guidance.
Always end with: {DIVINE_BLESSING}
User query: {{query}}
""",
    "action": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are the Action Engine. You detect the user's intention and perform tasks like creating documents, scheduling, sending emails, etc.
You respond with a clear action plan and optionally generate the required output (e.g., drafted email, calendar event JSON).
Always end with: {DIVINE_BLESSING}
User request: {{query}}
""",
    "gst_agent": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are Lord Kubera. Specialise in GST compliance, filing, returns, and registration.
Always end with: {DIVINE_BLESSING}
Query: {{query}}
""",
    "income_tax_agent": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are Goddess Lakshmi. Specialise in Income Tax Act, ITR filing, TDS, and tax planning.
Always end with: {DIVINE_BLESSING}
Query: {{query}}
""",
    "incorporation_agent": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are Lord Brahma. Specialise in company incorporation, ROC compliance, and drafting MOA/AOA.
Always end with: {DIVINE_BLESSING}
Query: {{query}}
""",
    "firm_registration_agent": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are Lord Vishnu. Specialise in partnership firm registration and LLP compliance.
Always end with: {DIVINE_BLESSING}
Query: {{query}}
""",
    "audit_agent": f"""{DIVINE_PREFACE}{DIVINE_SALUTATION}
You are Lord Yama. Specialise in statutory audit, internal audit, tax audit, and CARO reporting.
Always end with: {DIVINE_BLESSING}
Query: {{query}}
"""
}

LANG_MAP = {
    "en": "English","es": "Spanish","fr": "French","de": "German","pt": "Portuguese","it": "Italian",
    "nl": "Dutch","ru": "Russian","sv": "Swedish","pl": "Polish","tr": "Turkish",
    "hi": "Hindi","bn": "Bengali","sa": "Sanskrit","ar": "Arabic",
    "zh": "Chinese","ja": "Japanese","ko": "Korean","th": "Thai","vi": "Vietnamese",
    "id": "Indonesian","ms": "Malay","he": "Hebrew","el": "Greek"
}

def detect_action(query: str) -> Optional[Dict]:
    q = query.lower()
    if q.startswith(("create ", "draft ", "write ")):
        return {"type": "create_document", "verb": q.split()[0]}
    if q.startswith(("schedule ", "book ", "set up ")):
        return {"type": "schedule_event"}
    if q.startswith(("send ", "email ")):
        return {"type": "send_email"}
    if q.startswith(("find ", "search ", "look up ")):
        return {"type": "web_search"}
    if q.startswith(("calculate ", "compute ")):
        return {"type": "calculate"}
    return None

def route_agent(query: str, agent_id: str = "agent_001", oracle_mode: bool = False) -> str:
    q = query.lower()
    if oracle_mode:
        return "oracle"
    if "what is lexsarthi" in q or "who are you" in q or "tell me about yourself" in q:
        return "about_lexsarthi"
    if "gst" in q or "goods and services tax" in q or "gstr" in q:
        return "gst_agent"
    if "income tax" in q or "itr" in q or "tax return" in q or "tds" in q:
        return "income_tax_agent"
    if "incorporate" in q or "company formation" in q or "private limited" in q or "opc" in q or "llp" in q:
        return "incorporation_agent"
    if "firm" in q or "partnership" in q or "registration of firm" in q:
        return "firm_registration_agent"
    if "audit" in q or "statutory audit" in q or "internal audit" in q or "tax audit" in q:
        return "audit_agent"
    if "contract" in q or "agreement" in q or "review" in q:
        return "contract_review"
    if "case" in q or "judgment" in q or "research" in q:
        return "legal_research"
    if "draft" in q or "create" in q or "prepare" in q:
        return "drafting"
    if "due diligence" in q or "compliance" in q:
        return "due_diligence"
    action = detect_action(query)
    if action:
        return "action"
    return "general"

async def run_swarm(query: str, model: str, lang: str = "en") -> str:
    logger.info("Swarm initiated for query: %s", query[:100])
    research_prompt = f"{DIVINE_PREFACE}{DIVINE_SALUTATION}\nYou are Lord Hanuman. Find statutes and case laws for: {query}"
    research_response = await execute_ai_raw(research_prompt, query, model, lang)
    draft_prompt = f"{DIVINE_PREFACE}{DIVINE_SALUTATION}\nYou are Goddess Saraswati. Draft a legal document based on:\n{research_response[:1000]}\n\nOriginal: {query}"
    draft_response = await execute_ai_raw(draft_prompt, query, model, lang)
    review_prompt = f"{DIVINE_PREFACE}{DIVINE_SALUTATION}\nYou are Lord Kartikeya. Review this draft for risks:\n{draft_response[:1000]}"
    review_response = await execute_ai_raw(review_prompt, query, model, lang)
    final = f"📜 **RESEARCH (Hanuman):**\n{research_response}\n\n📝 **DRAFT (Saraswati):**\n{draft_response}\n\n✅ **REVIEW (Kartikeya):**\n{review_response}\n\n{DIVINE_BLESSING}"
    return final

async def execute_ai_raw(system_prompt: str, query: str, model: str, lang: str) -> str:
    if model.startswith("llama") and groq_client:
        try:
            response = groq_client.chat.completions.create(
                model=model,
                messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": query}],
                temperature=0.3,
                max_tokens=2048,
            )
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Groq error: {e}")
    return search_local_knowledge(query, has_file=False)

async def execute_ai(query: str, model: str, agent_type: str, agent_name: str, lang: str = "en", has_file: bool = False) -> str:
    # If swarm triggered
    if agent_type == "due_diligence" and "swarm" in query.lower():
        return await run_swarm(query, model, lang)

    # If action agent
    if agent_type == "action":
        action = detect_action(query)
        base = BASE_AGENT_PROMPTS["action"]
        prompt = base.format(query=query)
        lang_instruction = f"IMPORTANT: Respond in {LANG_MAP.get(lang, 'English')} language. Use appropriate script."
        system_prompt = f"{prompt}\n\n{lang_instruction}"
        if model.startswith("llama") and groq_client:
            try:
                response = groq_client.chat.completions.create(
                    model=model,
                    messages=[{"role": "system", "content": f"You are {agent_name}. {system_prompt}"}, {"role": "user", "content": query}],
                    temperature=0.3,
                    max_tokens=2048,
                )
                return response.choices[0].message.content
            except Exception as e:
                logger.error(f"Groq error: {e}")
        return f"I understand you want to perform an action: {action['type'] if action else 'unknown'}. I will process your request shortly."

    # Standard agents
    if agent_type == "oracle":
        prompt = ORACLE_PROMPT.format(query=query)
    else:
        base = BASE_AGENT_PROMPTS.get(agent_type, BASE_AGENT_PROMPTS["general"])
        prompt = base.format(query=query)
    
    lang_instruction = f"IMPORTANT: Respond in {LANG_MAP.get(lang, 'English')} language. Use appropriate script."
    system_prompt = f"{prompt}\n\n{lang_instruction}"
    
    # Try AI providers
    if model.startswith("llama") and groq_client:
        try:
            response = groq_client.chat.completions.create(
                model=model,
                messages=[{"role": "system", "content": f"You are {agent_name}. {system_prompt}"}, {"role": "user", "content": query}],
                temperature=0.3,
                max_tokens=4096,
            )
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Groq error: {e}")
    # Fallback – pass has_file flag to avoid nonsense library search on file uploads
    return search_local_knowledge(query, has_file=has_file)

# ─── API ENDPOINTS ──────────────────────────────────────────────
@app.get("/health")
async def health():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.post("/auth/login", response_model=Token)
async def login(user_login: UserLogin):
    logger.info(f"Login: {user_login.username}")
    user = await database.fetch_one(users.select().where(users.c.username == user_login.username))
    if not user:
        user = await database.fetch_one(users.select().where(users.c.email == user_login.username.lower()))
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    user = dict(user)
    if not verify_password(user_login.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    token = create_access_token({"sub": str(user["id"])})
    return {
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "id": user["id"],
            "username": user["username"],
            "email": user["email"],
            "full_name": user["full_name"],
            "tier": user["tier"],
            "is_premium": user["is_premium"],
            "api_key": user.get("api_key")
        }
    }

@app.post("/auth/register")
async def register(user: UserCreate):
    existing = await database.fetch_one(users.select().where((users.c.username == user.username) | (users.c.email == user.email.lower())))
    if existing:
        raise HTTPException(status_code=400, detail="User already exists")
    hashed = hash_password(user.password)
    api_key = ''.join(random.choices(string.ascii_letters + string.digits, k=32))
    stmt = users.insert().values(
        username=user.username,
        email=user.email.lower(),
        password_hash=hashed,
        full_name=user.full_name,
        tier="free",
        api_key=api_key,
        memory=[]
    ).returning(users.c.id)
    user_id = await database.fetch_val(stmt)
    token = create_access_token({"sub": str(user_id)})
    return {"access_token": token, "token_type": "bearer", "user": {"id": user_id, "username": user.username, "api_key": api_key}}

@app.get("/auth/me")
async def get_me(current_user: dict = Depends(get_current_user)):
    return current_user

@app.get("/lifetime-count")
async def lifetime_count():
    count = await database.fetch_val(select(func.count()).select_from(users).where(users.c.tier == "lifetime")) or 0
    return {"count": count, "limit": 1000, "remaining": max(0, 1000 - count)}

@app.get("/my-usage")
async def my_usage(current_user: dict = Depends(get_current_user)):
    total = await database.fetch_val(select(func.count()).select_from(queries).where(queries.c.user_id == current_user["id"])) or 0
    today = await database.fetch_val(select(func.count()).select_from(queries).where(
        queries.c.user_id == current_user["id"],
        func.date(queries.c.created_at) == func.current_date()
    )) or 0
    return {"total_queries": total, "queries_today": today}

@app.post("/ask")
@limiter.limit("30/minute")
async def ask(
    request: Request,
    query: str = Form(...),
    files: Optional[UploadFile] = File(None),
    search_web: str = Form("off"),
    model: str = Form("llama-3.3-70b-versatile"),
    agent_id: str = Form("agent_001"),
    lang: str = Form("en"),
    oracle_mode: str = Form("false"),
    current_user: dict = Depends(get_current_user)
):
    if not await check_query_limit(current_user):
        raise HTTPException(status_code=429, detail="Free limit reached.")
    
    combined_query = query
    has_file = False
    if files:
        has_file = True
        try:
            file_text = await process_file(files)
            if not file_text or len(file_text.strip()) < 20:
                raise HTTPException(status_code=400, detail="File is empty or contains no readable text.")
            combined_query = f"{query}\n\n--- Document Content ---\n{file_text}"
        except HTTPException as he:
            raise he
        except Exception as e:
            logger.error(f"File error: {e}")
            raise HTTPException(status_code=400, detail=f"File processing failed: {str(e)}")
    
    # --- NEW: Check query length to prevent AI failure ---
    if len(combined_query) > 20000:  # roughly 5000 tokens
        if has_file:
            raise HTTPException(
                status_code=413, 
                detail="Document is too large for real-time review. Please upload a smaller document (under 10 pages) or use the Bulk Upload feature for large batches."
            )
    
    if search_web.lower() in ("on", "yes"):
        combined_query += "\n\n--- Web Search Enabled ---"
    
    await increment_query(current_user["id"])
    
    memory = await get_user_memory(current_user["id"])
    context_prompt = build_context_prompt(memory)
    if context_prompt:
        combined_query = context_prompt + combined_query
    
    oracle = oracle_mode.lower() == "true"
    agent_type = route_agent(combined_query, agent_id, oracle)
    agent_name = next((a["name"] for a in DIVINE_AGENTS if a["id"] == agent_id), "General Counsel")
    
    # Pass has_file to execute_ai to prevent fallback nonsense
    response_text = await execute_ai(combined_query, model, agent_type, agent_name, lang, has_file)
    
    await update_user_memory(current_user["id"], query, response_text)
    
    expires_at = datetime.now() + timedelta(hours=24)
    await database.execute(
        queries.insert().values(
            user_id=current_user["id"],
            query=combined_query,
            response=response_text,
            metadata={"agent": agent_id, "model": model, "file": has_file, "agent_type": agent_type, "lang": lang, "oracle": oracle},
            expires_at=expires_at
        )
    )
    return {"response": response_text, "model": model, "agent_used": agent_id}

# ─── ACTION ENDPOINT ────────────────────────────────────────────
@app.post("/action")
async def execute_action(action_req: ActionRequest, current_user: dict = Depends(get_current_user)):
    action_type = action_req.action
    data = action_req.data
    await database.execute(
        actions.insert().values(
            user_id=current_user["id"],
            action_type=action_type,
            action_data=data,
            result={"status": "mocked", "message": f"Action '{action_type}' received."}
        )
    )
    return {"status": "ok", "action": action_type, "data": data}

# ─── BULK UPLOAD ────────────────────────────────────────────────
@app.post("/bulk-upload")
async def bulk_upload(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    query: str = Form(...),
    model: str = Form("llama-3.3-70b-versatile"),
    agent_id: str = Form("agent_001"),
    lang: str = Form("en"),
    current_user: dict = Depends(get_current_user)
):
    if current_user["tier"] not in ("premium", "enterprise", "lifetime"):
        raise HTTPException(status_code=403, detail="Bulk upload requires Premium or Enterprise.")
    job_id = str(uuid.uuid4())
    total_files = len(files)
    await database.execute(
        bulk_jobs.insert().values(
            user_id=current_user["id"],
            job_id=job_id,
            total_files=total_files,
            status="processing",
            expires_at=datetime.now() + timedelta(days=7)
        )
    )
    background_tasks.add_task(process_bulk_job, job_id, files, query, model, agent_id, lang, current_user["id"])
    return {"job_id": job_id, "status": "processing", "total_files": total_files}

async def process_bulk_job(job_id: str, files: List[UploadFile], query: str, model: str, agent_id: str, lang: str, user_id: int):
    results = []
    processed = 0
    for file in files:
        try:
            file_text = await process_file(file)
            combined_query = f"{query}\n\n--- Document Content ---\n{file_text}"
            agent_type = route_agent(combined_query, agent_id)
            agent_name = next((a["name"] for a in DIVINE_AGENTS if a["id"] == agent_id), "General Counsel")
            response = await execute_ai(combined_query, model, agent_type, agent_name, lang, has_file=True)
            results.append({"filename": file.filename, "response": response})
        except Exception as e:
            results.append({"filename": file.filename, "error": str(e)})
        processed += 1
        await database.execute(
            bulk_jobs.update().where(bulk_jobs.c.job_id == job_id).values(
                processed_files=processed,
                status="processing"
            )
        )
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Filename", "Response"])
    for r in results:
        writer.writerow([r.get("filename"), r.get("response", r.get("error", "Failed"))])
    csv_data = output.getvalue()
    await database.execute(
        bulk_jobs.update().where(bulk_jobs.c.job_id == job_id).values(
            status="completed",
            result_url=f"job_{job_id}.csv"
        )
    )

@app.get("/bulk-result/{job_id}")
async def get_bulk_result(job_id: str, current_user: dict = Depends(get_current_user)):
    job = await database.fetch_one(bulk_jobs.select().where(bulk_jobs.c.job_id == job_id))
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job["user_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="Access denied")
    if job["status"] != "completed":
        return {"status": job["status"], "processed": job["processed_files"], "total": job["total_files"]}
    return {"status": "completed", "result_url": job["result_url"]}

# ─── RAZORPAY ──────────────────────────────────────────────────
razorpay_client = razorpay.Client(auth=(RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET)) if RAZORPAY_KEY_ID else None

@app.post("/create-order")
async def create_order(payment: PaymentCreate, current_user: dict = Depends(get_current_user)):
    if not razorpay_client:
        raise HTTPException(status_code=501, detail="Payments not configured")
    amount_map = {"premium": 10200, "enterprise": 101100, "lifetime": 200}
    amount = amount_map.get(payment.tier, 10200)
    order = razorpay_client.order.create({"amount": amount, "currency": "INR", "payment_capture": 1})
    await database.execute(
        payments.insert().values(
            user_id=current_user["id"],
            razorpay_order_id=order["id"],
            amount=amount/100,
            tier=payment.tier,
            status="created"
        )
    )
    return {"order_id": order["id"], "amount": amount, "razorpay_key": RAZORPAY_KEY_ID}

@app.post("/verify-payment")
async def verify_payment(
    razorpay_order_id: str = Form(...),
    razorpay_payment_id: str = Form(...),
    razorpay_signature: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    if not razorpay_client:
        raise HTTPException(status_code=501, detail="Payments not configured")
    try:
        razorpay_client.utility.verify_payment_signature({
            "razorpay_order_id": razorpay_order_id,
            "razorpay_payment_id": razorpay_payment_id,
            "razorpay_signature": razorpay_signature
        })
        payment = await database.fetch_one(payments.select().where(payments.c.razorpay_order_id == razorpay_order_id))
        tier = payment["tier"]
        await database.execute(
            users.update().where(users.c.id == current_user["id"]).values(tier=tier, is_premium=True)
        )
        await database.execute(
            payments.update().where(payments.c.razorpay_order_id == razorpay_order_id).values(status="paid")
        )
        return {"status": "success", "tier": tier}
    except Exception as e:
        logger.error(f"Payment verification failed: {e}")
        raise HTTPException(status_code=400, detail="Verification failed")

# ─── STATIC FILES ──────────────────────────────────────────────
app.mount("/", StaticFiles(directory="static", html=True), name="static")

if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=7860, reload=False)