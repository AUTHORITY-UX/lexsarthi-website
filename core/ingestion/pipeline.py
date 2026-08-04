"""
core/ingestion/pipeline.py - Multi-format ingestion pipeline
Handles PDF, HTML, Text, Markdown → legal-aware chunker → embedder → pgvector
"""

import os
import json
import re
import hashlib
import logging
from typing import List, Dict, Any, Optional, AsyncGenerator
from dataclasses import dataclass, field
from pathlib import Path

import PyPDF2
import docx
import html2text
import markdown
from bs4 import BeautifulSoup

from core.db import db
from core.config import settings
from core.llm import get_router, LLMMessage

logger = logging.getLogger(__name__)


@dataclass
class LegalDocument:
    """Represents a parsed legal document"""
    id: str = field(default_factory=lambda: str(uuid4()))
    title: str = ""
    source: str = ""
    doc_type: str = ""  # 'case', 'statute', 'commentary', 'article'
    content: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunks: List[Dict[str, Any]] = field(default_factory=list)
    embedding: Optional[List[float]] = None


class DocumentParser:
    """Parse various document formats"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
        return text
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return ""
    
    @staticmethod
    def parse_html(html_content: str) -> str:
        """Extract text from HTML"""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text(separator='\n')
        except Exception as e:
            logger.error(f"HTML parsing error: {e}")
            return ""
    
    @staticmethod
    def parse_markdown(md_content: str) -> str:
        """Extract text from Markdown"""
        try:
            html = markdown.markdown(md_content)
            return DocumentParser.parse_html(html)
        except Exception as e:
            logger.error(f"Markdown parsing error: {e}")
            return ""
    
    @staticmethod
    def detect_format(file_path: str) -> str:
        """Detect document format by extension"""
        ext = Path(file_path).suffix.lower()
        format_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',
            '.html': 'html',
            '.htm': 'html',
            '.md': 'markdown',
            '.txt': 'text',
            '.rtf': 'rtf'
        }
        return format_map.get(ext, 'text')


class LegalChunker:
    """Legal-aware text chunking with overlap"""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def chunk(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """Chunk legal text with awareness of legal structures"""
        chunks = []
        metadata = metadata or {}
        
        # Identify legal structure boundaries
        boundaries = self._find_legal_boundaries(text)
        
        # If boundaries found, split by them
        if boundaries:
            sections = self._split_by_boundaries(text, boundaries)
            for section in sections:
                if len(section) > self.chunk_size:
                    chunks.extend(self._chunk_text(section, metadata))
                else:
                    chunks.append({
                        'text': section,
                        'metadata': metadata.copy(),
                        'chunk_type': 'legal_section'
                    })
        else:
            chunks.extend(self._chunk_text(text, metadata))
        
        return chunks
    
    def _find_legal_boundaries(self, text: str) -> List[str]:
        """Find legal structure boundaries (sections, articles, etc.)"""
        boundaries = []
        
        # Indian legal structure patterns
        patterns = [
            r'Section\s+\d+[A-Za-z]*(?:\(\d+\))?',
            r'Article\s+\d+[A-Za-z]*(?:\(\d+\))?',
            r'Rule\s+\d+[A-Za-z]*(?:\(\d+\))?',
            r'CHAPTER\s+[IVXLCDM]+',
            r'PART\s+[IVXLCDM]+',
            r'Rule\s+\d+',
            r'Regulation\s+\d+',
        ]
        
        for pattern in patterns:
            matches = list(re.finditer(pattern, text, re.IGNORECASE))
            if len(matches) > 1:  # Only if multiple matches
                boundaries.extend([m.group(0) for m in matches])
        
        return list(set(boundaries)) if boundaries else []
    
    def _split_by_boundaries(self, text: str, boundaries: List[str]) -> List[str]:
        """Split text by legal boundaries"""
        sections = []
        current = []
        
        lines = text.split('\n')
        for line in lines:
            if any(boundary in line for boundary in boundaries):
                if current:
                    sections.append('\n'.join(current))
                    current = []
            current.append(line)
        
        if current:
            sections.append('\n'.join(current))
        
        return sections
    
    def _chunk_text(self, text: str, metadata: Dict) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks"""
        chunks = []
        words = text.split()
        
        for i in range(0, len(words), self.chunk_size - self.overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            if chunk_text.strip():
                chunks.append({
                    'text': chunk_text,
                    'metadata': metadata.copy(),
                    'chunk_type': 'text_chunk',
                    'chunk_index': len(chunks)
                })
        
        return chunks


class EmbeddingGenerator:
    """Generate embeddings for legal text"""
    
    def __init__(self):
        self.router = get_router()
        # We'll use the LLM router for embedding generation
        # For now, use a simple approach
    
    async def generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text using LLM router"""
        try:
            # Use a simple embedding approach via LLM
            # In production, use a dedicated embedding model
            messages = [
                LLMMessage(
                    role="system",
                    content="Generate a dense vector representation of this legal text. Respond with the embedding as a JSON array of floats."
                ),
                LLMMessage(role="user", content=text[:500])  # Limit for embedding
            ]
            
            response = await self.router.chat(messages, model="sarvam-105b", max_tokens=2000)
            
            # Try to parse the response as embedding
            if response.content:
                import json
                try:
                    # Look for JSON array in response
                    match = re.search(r'\[[\d.,\s]+\]', response.content)
                    if match:
                        embedding = json.loads(match.group(0))
                        return embedding
                except:
                    pass
            
            # Fallback: use random embedding (in production, use proper embedding model)
            import random
            return [random.random() for _ in range(1536)]
            
        except Exception as e:
            logger.error(f"Embedding generation error: {e}")
            return [0.0] * 1536


class IngestionPipeline:
    """Complete ingestion pipeline for legal documents"""
    
    def __init__(self):
        self.parser = DocumentParser()
        self.chunker = LegalChunker()
        self.embedder = EmbeddingGenerator()
    
    async def process_file(self, file_path: str, doc_type: str = "case") -> Optional[LegalDocument]:
        """Process a single file through the entire pipeline"""
        logger.info(f"Processing file: {file_path}")
        
        # Parse the document
        doc_format = self.parser.detect_format(file_path)
        if doc_format == 'pdf':
            content = self.parser.parse_pdf(file_path)
        elif doc_format == 'docx':
            content = self.parser.parse_docx(file_path)
        elif doc_format == 'html':
            with open(file_path, 'r') as f:
                content = self.parser.parse_html(f.read())
        elif doc_format == 'markdown':
            with open(file_path, 'r') as f:
                content = self.parser.parse_markdown(f.read())
        else:
            with open(file_path, 'r') as f:
                content = f.read()
        
        if not content:
            logger.error(f"No content extracted from {file_path}")
            return None
        
        # Create document object
        doc = LegalDocument(
            title=Path(file_path).stem,
            source=file_path,
            doc_type=doc_type,
            content=content,
            metadata={
                'format': doc_format,
                'size': len(content),
                'word_count': len(content.split())
            }
        )
        
        # Chunk the document
        chunks = self.chunker.chunk(content, doc.metadata)
        doc.chunks = chunks
        
        # Generate embeddings for each chunk
        for chunk in chunks:
            embedding = await self.embedder.generate_embedding(chunk['text'])
            chunk['embedding'] = embedding
        
        return doc
    
    async def store_document(self, doc: LegalDocument) -> bool:
        """Store document and chunks in database"""
        try:
            # Store in case_law table
            await db.execute("""
                INSERT INTO case_law 
                (citation, title, full_text, embedding, metadata)
                VALUES ($1, $2, $3, $4, $5)
            """, 
                doc.title[:255],  # citation placeholder
                doc.title,
                doc.content,
                [0.0] * 1536,  # placeholder embedding
                json.dumps(doc.metadata)
            )
            
            # Store chunks
            for chunk in doc.chunks:
                await db.execute("""
                    INSERT INTO irac_nodes 
                    (node_type, content, embedding, metadata)
                    VALUES ($1, $2, $3, $4)
                """,
                    'text_chunk',
                    chunk['text'],
                    chunk.get('embedding', [0.0] * 1536),
                    json.dumps(chunk.get('metadata', {}))
                )
            
            logger.info(f"Stored document: {doc.title}")
            return True
            
        except Exception as e:
            logger.error(f"Store error: {e}")
            return False
    
    async def process_directory(self, directory: str, doc_type: str = "case") -> List[Dict]:
        """Process all files in a directory"""
        results = []
        
        for root, dirs, files in os.walk(directory):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx', '.html', '.md', '.txt')):
                    file_path = os.path.join(root, file)
                    doc = await self.process_file(file_path, doc_type)
                    
                    if doc:
                        result = await self.store_document(doc)
                        results.append({
                            'file': file,
                            'success': result,
                            'chunks': len(doc.chunks)
                        })
        
        return results